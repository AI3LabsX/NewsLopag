import asyncio
import os
from urllib.parse import urlparse

import feedparser
import pytesseract
import requests
from PIL import Image
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup
from dateutil import parser
from docx import Document
from openai import AsyncOpenAI
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, ForceReply
from telegram.ext import (
    CallbackQueryHandler,
    MessageHandler,
    filters,
)
from telegram.ext import CommandHandler, CallbackContext

from tg.handlers.db_tools import get_topics, add_topic_to_db, delete_topic_from_db, delete_rss_feed_from_db, \
    get_rss_feeds, add_rss_feed_to_db
from tg.handlers.decarators import exponential_backoff_async
from tg.utils.filters import is_admin_filter
from tg.utils.logger import logger


async def start_cmd_from_admin(update: Update, context: CallbackContext) -> None:
    """Handles command /start from the user"""
    await update.message.reply_text(
        f"Welcome. Type /settings to set up the bot"
    )
    context.user_data["conversation_state"] = "idle"


async def start_cmd_from_user(update: Update, context: CallbackContext) -> None:
    """Handles command /start from the user"""
    await update.message.reply_text(
        f"Welcome. Type /settings to set up the bot"
    )
    context.user_data["conversation_state"] = "idle"


async def settings_command(update: Update, context: CallbackContext):
    keyboard = [
        [InlineKeyboardButton("1 - Choose Topic", callback_data='choose_topic')],
        [InlineKeyboardButton("2 - Add/Delete RSS feed", callback_data='add_delete_rss')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    sent_message = await update.message.reply_text('Choose one of the settings:', reply_markup=reply_markup)
    context.user_data['original_message_id'] = sent_message.message_id


async def callback_handler(update: Update, context: CallbackContext):
    query = update.callback_query
    callback_data = query.data
    user_id = query.from_user.id

    try:
        await query.answer()
        topics = await get_topics(user_id)
        message_text = "Here is you Settings:\n\nYour Topics: " + (', '.join(topics) if topics else "No topics added.")

        if callback_data.startswith("choose_topic"):
            keyboard = [
                [InlineKeyboardButton("Add Topic", callback_data='add_topic')],
                [InlineKeyboardButton("Delete Topic", callback_data='delete_topic')],
                [InlineKeyboardButton("Confirm Changes", callback_data='confirm_changes')]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(text=message_text, reply_markup=reply_markup)

        elif callback_data.startswith("add_topic"):
            keyboard = [
                [InlineKeyboardButton("Confirm Changes", callback_data='confirm_changes')]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            context.user_data['adding_topic'] = True
            await query.edit_message_text(text=message_text + "\n\nPlease type the topic you want to add and send.",
                                          reply_markup=reply_markup)

        elif callback_data.startswith("delete_topic"):
            keyboard = [[InlineKeyboardButton(topic, callback_data=f"delete_top_{topic}")] for topic in topics]
            keyboard.append([InlineKeyboardButton("Back to Settings", callback_data='choose_topic')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(text="Select a topic to delete:", reply_markup=reply_markup)

        elif callback_data.startswith("delete_"):
            selected = callback_data.split("_")[-1]
            if "_top_" in callback_data:
                keyboard = [
                    [InlineKeyboardButton("Add Topic", callback_data='add_topic')],
                    [InlineKeyboardButton("Delete Topic", callback_data='delete_topic')],
                    [InlineKeyboardButton("Confirm Changes", callback_data='confirm_changes')]
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)
                await delete_topic_from_db(user_id, selected)
                updated_topics = await get_topics(user_id)
                await query.edit_message_text(
                    text="Settings updated. \n\nYour Topics: " + (
                        ', '.join(updated_topics) if updated_topics else "No topics added."), reply_markup=reply_markup)

            else:
                keyboard = [
                    [InlineKeyboardButton("Add RSS Feed", callback_data='add_rss')],
                    [InlineKeyboardButton("Delete RSS Feed", callback_data='deletes_rss')],
                    [InlineKeyboardButton("Back to Settings", callback_data='choose_topic')]
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)
                await delete_rss_feed_from_db(user_id, selected)
                updated_rss = await get_rss_feeds(user_id)
                await query.edit_message_text(
                    text="Here is you Settings:\n\nYour RSS Feeds: " + (
                        ', '.join(
                            [urlparse(str(rss)).hostname for rss in updated_rss]) if updated_rss else "No RSS added."),
                    reply_markup=reply_markup)

        elif callback_data.startswith("add_delete_rss"):
            rss_feeds = await get_rss_feeds(user_id)
            message_text = "Your RSS Feeds: " + (', '.join(
                [urlparse(str(feed)).hostname for feed in rss_feeds]) if rss_feeds else "No RSS feeds added.")
            keyboard = [
                [InlineKeyboardButton("Add RSS Feed", callback_data='add_rss')],
                [InlineKeyboardButton("Delete RSS Feed", callback_data='deletes_rss')],
                [InlineKeyboardButton("Back to Settings", callback_data='choose_topic')]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(text=message_text, reply_markup=reply_markup)

        elif callback_data.startswith("add_rss"):
            context.user_data['adding_rss'] = True
            await query.edit_message_text(text="Please type the RSS feed URL you want to add and send.",
                                          reply_markup=InlineKeyboardMarkup([]))

        elif callback_data.startswith("deletes_rss"):
            rss_feeds = await get_rss_feeds(user_id)
            keyboard = [[InlineKeyboardButton(urlparse(str(feed)).hostname, callback_data=f"delete_rss_{feed}")] for
                        feed in rss_feeds]
            keyboard.append([InlineKeyboardButton("Back to RSS Feeds", callback_data='add_delete_rss')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(text="Select an RSS feed to delete:", reply_markup=reply_markup)

        elif callback_data.startswith("confirm_changes"):
            await query.edit_message_text(text=f"Settings updated. Your current topics: {', '.join(topics)}. Press "
                                               f"/settings to set up rss feeds. Or just wait for news")
    except Exception as e:
        await query.edit_message_text(f"An error occurred: {e}")


def is_valid_url(url):
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except ValueError:
        return False


# TODO: Test it all : https://chat.openai.com/c/84c7fc15-d61b-4a6e-8cfb-78910aa745fa


async def fetch_rss_feed(feed_url):
    """Fetch and parse RSS feed."""
    logger.info(f"Fetching RSS feed from URL: {feed_url}")
    return feedparser.parse(feed_url)


async def article_check(article, topic):
    client = AsyncOpenAI(
        # This is the default and can be omitted
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    data = {
        "model": "gpt-4-1106-preview",
        "messages": [
            {
                "role": "system",
                "content": f"You only job is to evaluate article title and first 500 words on context, if it is fits "
                           f"with one of the topics chosen by user. If article fits under the description of one of "
                           f"the topic, output only True, otherwise False. PLease insure that you are doing it "
                           f"correctly, topic can be valid, when using exact keyword, talking about topic iin content "
                           f"or has any affect of topic directly or passively Only provide one word as answer: "
                           f"True/False.",
            },
            {
                "role": "user",
                "content": f"Article title {article['title']}\nArticle Content: {article['content']}.\nUser topics {topic}",
            },
        ],
        "temperature": 0,
        "max_tokens": 1500,
        "top_p": 0.4,
        "frequency_penalty": 1.5,
        "presence_penalty": 1,
    }
    response = await client.chat.completions.create(**data)

    answer = response.choices[0].message.content
    logger.info(answer)
    return answer.strip().lower() == 'true'


def extract_largest_text_block(soup: BeautifulSoup) -> str:
    paragraphs = soup.find_all('p')
    largest_block = ""
    current_block = ""
    for paragraph in paragraphs:
        if len(paragraph.text) > 50:
            current_block += paragraph.text + "\n"
        else:
            if len(current_block) > len(largest_block):
                largest_block = current_block
            current_block = ""
    return largest_block.strip()


def get_final_url(url: str) -> str:
    if "news.google.com" in url:
        response = requests.get(url, allow_redirects=True)
        return response.url
    return url


async def summary_article(article):
    client = AsyncOpenAI(
        # This is the default and can be omitted
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    data = {
        "model": "gpt-3.5-turbo-16k",
        "messages": [
            {
                "role": "system",
                "content": f"You are AI news bot. Create short summary of the article provided by user. Insure to "
                           f"cover all main points, while providing cosine response! Please be short, "
                           f"try to summarize the text under 50 words.• Concise Summary: Create a very brief summary "
                           f"for each article, ideally in one or two sentences, capturing the essence of the "
                           f"article.• Highlighted Stats: List all crucial statistics and numbers in a highlighted "
                           f"format the summary for emphasis and clarity.• Final Output:• Compile a collection of "
                           f"short, structured highlights for each source.• Combine highlights from all sources, "
                           f"ensuring no repetition and that each news item is represented once in the concise form "
                           f"possible without leaving any names and stats shared in article.",
            },
            {
                "role": "user",
                "content": f"Article Content: {article['content']}",
            },
        ],
        "temperature": 0,
        "max_tokens": 2000,
        "top_p": 0.4,
        "frequency_penalty": 1.5,
        "presence_penalty": 1,
    }
    response = await client.chat.completions.create(**data)
    return response.choices[0].message.content


@exponential_backoff_async(max_iterations=3, exponent_limit=2)
async def monitor_rss_feeds_for_user(user_id, context, update):
    logger.debug(f"Starting RSS feed monitoring for user: {user_id}")
    chat_id = update.effective_chat.id

    while True:
        try:
            rss_feeds = await get_rss_feeds(user_id)
            logger.debug(f"Retrieved RSS feeds for user {user_id}: {rss_feeds}")
            if not rss_feeds:
                logger.warning(f"No RSS feeds found for user: {user_id}")
                break

            if user_id not in context.user_data:
                context.user_data[user_id] = {'latest_article': {}}
            logger.debug(f"User data for {user_id}: {context.user_data[user_id]}")

            total_news = ""
            for rss_feed in rss_feeds:
                feed = await fetch_rss_feed(rss_feed)
                logger.debug(f"Fetched feed for {rss_feed}: {len(feed.entries)} entries found")
                if feed.entries:
                    latest_article = feed.entries[0]
                    latest_article_date = parser.parse(latest_article.published)
                    logger.debug(f"Latest article date for feed {rss_feed}: {latest_article_date}")

                    if (rss_feed not in context.user_data[user_id]['latest_article']) or \
                            (latest_article_date > context.user_data[user_id]['latest_article'][rss_feed]):
                        captions = await process_article(user_id, latest_article)
                        if captions:
                            total_news += captions
                        context.user_data[user_id]['latest_article'][rss_feed] = latest_article_date
                    else:
                        logger.debug(f"No new article to process for feed {rss_feed}")

            if total_news:
                await context.bot.send_message(chat_id=chat_id, text=total_news)
                logger.debug(f"Sent message to chat {chat_id}")
            else:
                logger.debug(f"No news to send for user {user_id}")

            await asyncio.sleep(300)  # Check feeds every 5 minutes
        except Exception as e:
            logger.error(f"Error in monitor_rss_feeds_for_user for user {user_id}: {e}")
        finally:
            if user_id not in context.bot_data.get('rss_users', {}):
                logger.info(f"Stopping RSS feed monitoring for user: {user_id}")
                break

# TODO: Here will be the logic for decision making if the article from feed corresponds with the topic. We can
#  use gpt to ask AI if article title + content
async def process_article(user_id, article):
    """Process a new article."""
    # Implement your logic here
    updated_topics = await get_topics(user_id)
    logger.info(f"User {user_id} topics are {updated_topics}")
    logger.info(f"Article Title {article.title}")

    title = article.title
    link = article.link
    final_link = get_final_url(link)
    html_code = requests.get(final_link)
    pub_date = parser.parse(article.published)
    soup = BeautifulSoup(html_code.text, 'html.parser')
    content = extract_largest_text_block(soup)
    image_div = soup.find('figure', {'class': 'article__lead__image'})
    image_url = image_div.find('img')['src'] if image_div else None
    articles = {
        "title": title,
        "link": link,
        "content": content,
        "pub_date": pub_date,
        "image": image_url
    }
    if_article = await article_check(articles, updated_topics)
    if if_article:
        summary = await summary_article(articles)
        captions = f"<b>Title</b>: {article.title}\n<b>Summary</b>: {summary.strip()}\n<a href='{article.link}'><b>Source</b></a>\n\n"
        logger.info(f"Article {article.title} was successfully sent ")
        return captions


async def news_on(update: Update, context: CallbackContext):
    user_id = update.effective_user.id
    if 'rss_users' not in context.bot_data:
        context.bot_data['rss_users'] = {}
    if user_id not in context.bot_data['rss_users']:
        context.bot_data['rss_users'][user_id] = asyncio.create_task(
            monitor_rss_feeds_for_user(user_id, context, update))
        logger.info(f"News monitoring activated for user: {user_id}")
        await update.message.reply_text("News monitoring activated.")
    else:
        logger.warning(f"Attempted to activate news monitoring for user {user_id}, but it's already active.")
        await update.message.reply_text("News monitoring is already active.")


async def news_off(update: Update, context: CallbackContext):
    user_id = update.effective_user.id
    if user_id in context.bot_data.get('rss_users', {}):
        task = context.bot_data['rss_users'].pop(user_id)
        task.cancel()
        logger.info(f"News monitoring deactivated for user: {user_id}")
        await update.message.reply_text("News monitoring deactivated.")
    else:
        logger.warning(f"Attempted to deactivate news monitoring for user {user_id}, but it was not active.")
        await update.message.reply_text("News monitoring is already deactivated.")


async def handle_text(update: Update, context: CallbackContext):
    user_id = update.message.from_user.id
    text = update.message.text
    original_message_id = context.user_data.get('original_message_id')
    chat_id = update.effective_chat.id

    if context.user_data.get('adding_topic'):
        if text.lower() != 'confirm':
            await add_topic_to_db(user_id, text)
            await update.message.delete()  # Delete the user's message to avoid clutter
            updated_topics = await get_topics(user_id)
            message_text = "Add more topics. Press Confirm, when you are done!:\n\nYour Topics: " + (
                ', '.join(updated_topics) if updated_topics else "No topics added.")
            reply_markup = InlineKeyboardMarkup([
                [InlineKeyboardButton("Confirm Changes", callback_data='confirm_changes')]
            ])
        else:
            context.user_data['adding_topic'] = False  # Clear the flag
            updated_topics = await get_topics(user_id)
            message_text = f"Settings updated. Your current topics: {', '.join(updated_topics) if updated_topics else 'None'}"
            reply_markup = InlineKeyboardMarkup([
                [InlineKeyboardButton("Add RSS Feed", callback_data='add_rss')],
                [InlineKeyboardButton("Delete RSS Feed", callback_data='deletes_rss')],
                [InlineKeyboardButton("Back to Settings", callback_data='choose_topic')]
            ])

        await context.bot.edit_message_text(chat_id=chat_id, message_id=original_message_id, text=message_text,
                                            reply_markup=reply_markup)

    elif context.user_data.get('adding_rss'):
        if text.lower() != 'confirm':
            if is_valid_url(text):
                await add_rss_feed_to_db(user_id, text)
                await update.message.delete()  # Delete the user's message
                updated_rss_feeds = await get_rss_feeds(user_id)
                message_text = "Your RSS Feeds: " + (', '.join([urlparse(str(feed)).hostname for feed in
                                                                updated_rss_feeds]) if updated_rss_feeds else "No RSS feeds added.")
            else:
                await update.message.delete()  # Delete the user's message
                message_text = "Invalid URL. Please enter a valid RSS feed URL."
        else:
            context.user_data['adding_rss'] = False  # Clear the flag
            updated_rss_feeds = await get_rss_feeds(user_id)
            message_text = "Settings updated. Your RSS feeds: " + (', '.join(
                [urlparse(str(feed)).hostname for feed in updated_rss_feeds]) if updated_rss_feeds else "None")

        reply_markup = InlineKeyboardMarkup([
            [InlineKeyboardButton("Add RSS Feed", callback_data='add_rss')],
            [InlineKeyboardButton("Delete RSS Feed", callback_data='deletes_rss')],
            [InlineKeyboardButton("Back to Settings", callback_data='choose_topic')]])
        await context.bot.edit_message_text(chat_id=chat_id, message_id=original_message_id, text=message_text,
                                            reply_markup=reply_markup)

    elif context.user_data.get('awaiting_language'):
        target_language = update.message.text
        extracted_text = context.user_data['extracted_text']

        # Perform translation
        translated_text = await translate_text(extracted_text, target_language)
        print(translated_text)
        # Create a new Document
        doc = Document()
        doc.add_heading('Translated Text', level=1)
        doc.add_paragraph(translated_text)

        # Save the document
        output_file = f"translated_{update.message.from_user.id}.docx"
        doc.save(output_file)

        # Send the document
        with open(output_file, 'rb') as file:
            await update.message.reply_document(document=file)

        # Delete the document and input file
        os.remove(output_file)
        if 'input_file_path' in context.user_data:
            os.remove(context.user_data['input_file_path'])
            del context.user_data['input_file_path']

        # Clean up other context data
        del context.user_data['extracted_text']
        del context.user_data['awaiting_language']


async def handle_file(update: Update, context: CallbackContext):
    if update.message.document:
        file = await update.message.document.get_file()
        file_extension = file.file_path.split('.')[-1].lower()

        # Download the file
        file_path = f"{file.file_id}.{file_extension}"
        await file.download_to_drive(custom_path=file_path)

        # Extract text based on file type
        extracted_text = ''
        if file_extension in ['jpg', 'jpeg', 'png']:
            extracted_text = extract_text_from_image(file_path)
        elif file_extension == 'pdf':
            extracted_text = extract_text_from_pdf(file_path)
        elif file_extension in ['docx']:
            extracted_text = extract_text_from_docx(file_path)

        # Store the extracted text and set a flag indicating we're waiting for a language code

        context.user_data['extracted_text'] = extracted_text
        context.user_data['awaiting_language'] = True
        context.user_data['input_file_path'] = file_path
        print(context.user_data['extracted_text'])
        # Ask the user for the language to translate to with ForceReply
        await update.message.reply_text(
            "Please reply with the language code (e.g., 'en' for English) to translate the text into.",
            reply_markup=ForceReply(selective=True)
        )


def extract_text_from_image(image_path):
    return pytesseract.image_to_string(Image.open(image_path))


def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ''
    for page in reader.pages:
        text += page.extract_text()
    return text


def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return '\n'.join([para.text for para in doc.paragraphs])


async def translate_text(text, language):
    client = AsyncOpenAI(
        # This is the default and can be omitted
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    text = text.replace('\n', ' ')
    data = {
        "model": "gpt-4-turbo-preview",
        "messages": [
            {
                "role": "system",
                "content": "PLease Translate the text into the langauge Specified by user! Please try to structure "
                           "the text and make it beautiful without changing context. PLease provide full word to word professional translation! Ensure that output text and input text have the same meaning. Please also fix all unicode characters"
                           "Insure that full text is provide as output to the user, no need for anything extra! For example suer can ask to translate text into Hindi, so provide good transaltion of full text!",
            },
            {
                "role": "user",
                "content": f"Text: {text}\nIn langauge: {language}",
            },
        ],
        "max_tokens": 4000,
        "temperature": 0,
        "top_p": 0.4,
        "frequency_penalty": 1.5,
        "presence_penalty": 1,
    }
    response = await client.chat.completions.create(**data)
    return response.choices[0].message.content


file_handler: MessageHandler = MessageHandler(filters.ALL, handle_file)

start_cmd_from_admin_handler: CommandHandler = CommandHandler(
    command="start", callback=start_cmd_from_admin, filters=is_admin_filter
)

start_cmd_from_user_handler: CommandHandler = CommandHandler(
    command="start", callback=start_cmd_from_user
)

settings_handler: CommandHandler = CommandHandler(command="settings", callback=settings_command)

button_callback_handler: CallbackQueryHandler = CallbackQueryHandler(
    callback_handler
)

add_topics_handler: MessageHandler = MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text)

news_on_handler: CommandHandler = CommandHandler('news_on', callback=news_on)
news_off_handler: CommandHandler = CommandHandler('news_off', news_off)
